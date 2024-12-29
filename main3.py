import streamlit as st
import pandas as pd
# from langchain_community.chat_models import ChatOpenAI
from langchain_openai import ChatOpenAI
# from langchain_core.runnables import RunnableSequence
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
# from langchain.chains import LLMChain
from dotenv import load_dotenv
import datetime, os
import threading
from io import StringIO, BytesIO
import csv
import sys
import docx

# Load environment variables from .env file
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Initialize OpenAI model
llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, model="gpt-4o-mini")

# Streamlit app title
st.title("AI SEO Content Generator")

# Create inputs for the landing page
st.header("Enter Business Details")

# This stops prompts from iterating over characters 
# instead of items, like sections or drafts.
error_ct = 0

business_name = st.text_input("Business Name", placeholder="Enter the business name")
business_description = st.text_area("Business Description", placeholder="Describe your business")
business_type = st.text_input("Business Type", placeholder="e.g., SaaS, E-commerce, Service")
niche = st.text_input("Niche", placeholder="e.g., Technology, Health, Education")
target_audience = st.text_input("Target Audience", placeholder="Describe your target audience")
features_page = st.text_input("Features Page", placeholder="https://www.yoursite.com/features")
money_page = st.text_input("Money Page", placeholder="https://www.yoursite.com/pricing")

keyword_focus = st.selectbox(
    "Keyword Focus", 
    ["Bottom of Funnel with high buying intent", 
     "Top of Funnel for building domain authority", 
     ]
)

topic = st.text_input("Topic", placeholder="Topic, e.g. churn for SaaS")
if not topic:
    topic = 'General topics based on above selections.'

# Define prompt templates
final_system_prompt = "You are a helpful Search Engine Optimization AI assistant."

foundation_template = PromptTemplate(
    input_variables=["business_name", "business_description", "business_type", "niche", "target_audience", "keyword_focus", "topic"],
    template="""
    Please provide a comprehensive SEO strategy report for the following business:
    - Business Name: {business_name}
    - Description: {business_description}
    - Type: {business_type}
    - Niche: {niche}
    - Target Audience: {target_audience}
    - Keyword Focus: {keyword_focus}
    - Topic: {topic}

    The report should include:
    1. Keyword research findings, including relevant keywords and their search volumes. Ensure these elements relate to the topic.
    2. Pillar page ideas.
    3. Website structure recommendations.
    """
)

keywords_template = PromptTemplate(
    input_variables=["foundation_content"],
    template="""
    {foundation_content}
    For the above list of keywords, treat them all as primary keywords. For each primary keyword, please find up to 10 secondary keywords that are related. Put the 
    results in a list, where the main headings are the primary keywords and the subheadings are the secondary keywords. 
    """
)

pillars_template = PromptTemplate(
    input_variables=["foundation_content"],
    template="""
    {foundation_content}
    For the above list of pillar page ideas, please create a new list. This new list will have headings that are the pillar page ideas. The subheadings will be sub-topics 
    associated with the pillar page idea. 
    """
)

questions_template = PromptTemplate(
    input_variables=["foundation_content"],
    template="""
    {foundation_content}
    Given the above pillar page ideas and sub-topics, and also taking into account all of the keywords mentioned before, what are some questions that my target audience might be 
    asking in search engines or on social media? Please list up to 20. Write the output from the perspective of my target audience. Group the questions under sub-headings that 
    are representative of the main idea behind the group of questions.
    """
)

pillar_extraction_template = PromptTemplate(
    input_variables=["foundation_content"],
    template="""
    From the below content, please extract only the pillar page ideas. Your output should simply be a string representation of a Python list where the values are the 
    pillar page ideas.

    {foundation_content}
    """
)

outline_template = PromptTemplate(
    input_variables=["pillar_page_idea", "foundation_content", "keywords", "questions"],
    template="""
    {foundation_content}
    {keywords}
    {questions}

    Keeping in mind all of the keywords and questions above, please generate an outline for the following pillar page idea: {pillar_page_idea}. Use the sub-topics associated 
    with that idea for the sections of this pillar page.

    The outline should follow this basic structure:

    - Introduction: Provide an overview of the topic and its importance to the target audience.
    - Section on SUB-TOPIC 1: Provide detailed information on SUB-TOPIC 1, including tips, strategies, and best practices.
    - Section on SUB-TOPIC 2: Provide detailed information on SUB-TOPIC 2, including tips, strategies, and best practices.
    etc.
    - Conclusion: Summarize the sub-topics and provide actionable takeaways for the target audience.

    Please provide as much detail as possible about the pillar page topic and the sub-topics to ensure that the outline is comprehensive. The output should be a formatted 
    outline. Please do not confirm the request; simply start the outline.
    """
)

firstdraft_template = PromptTemplate(
    input_variables=["outline", "business_name", "business_description", "features_page", "money_page"],
    template="""
    Business Name: {business_name}
    Description: {business_description}
    Features page: {features_page}
    Money page: {money_page}

    Please expand the following outline into a 1000 word blog post. Incorporate a casual, conversational tone, use humor where appropriate to make the writing more engaging, and make sure the message is clear and persuasive. Be sure to weave in sensible CTAs related to the service that this business provides.

    This draft should have a title in an h1 headings, sub headings that are H2, and if appropriate, more subheadings that are H3s. The output should be a formatted blog post which uses the previously discussed keywords and answers at least some of the previously discussed questions.

    outline:
    <outline>
    {outline}
    </outline>
    """
)

sectionextract_template = PromptTemplate(
    input_variables=["firstdraft"],
    template="""
    <firstdraft>
    {firstdraft}
    </firstdraft>

    Given the above first draft, please extract each meaningful and cohesive section. Your output should be a python list, where each list item is one of the sections.
    """
)

sectiondev_template = PromptTemplate(
    input_variables=["outline", "business_name", "business_description", "features_page", "money_page"],
    template="""
    Business Name: {business_name}
    Description: {business_description}
    Features page: {features_page}
    Money page: {money_page}

    Please expand the following outline into a 1000 word blog post. Incorporate a casual, conversational tone, use humor where appropriate to make the writing more engaging, and make sure the message is clear and persuasive. Be sure to weave in sensible CTAs related to the service that this business provides.

    This draft should have a title in an h1 headings, sub headings that are H2, and if appropriate, more subheadings that are H3s. The output should be a formatted blog post which uses the previously discussed keywords and answers at least some of the previously discussed questions.

    outline:
    <outline>
    {outline}
    </outline>
    """
)

# Export functions
def export_to_csv(data, filename):
    csv_data = StringIO()
    writer = csv.writer(csv_data)
    writer.writerow(["Section", "Content"])
    for section, content in data.items():
        writer.writerow([section, content])
    st.download_button(
        label="Download CSV",
        data=csv_data.getvalue(),
        file_name=filename,
        mime="text/csv"
    )

def export_to_word(data, filename):
    doc = docx.Document()
    for section, content in data.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    doc_data = BytesIO()
    doc.save(doc_data)
    doc_data.seek(0)  # Move the pointer to the start of the stream
    st.download_button(
        label="Download Word Document",
        data=doc_data.getvalue(),
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Some other helper functions for taking drafts to more developed articles
def extract_sections(draft_content):
    """Use LangChain to extract sections from a draft."""
    section_extract_prompt = PromptTemplate(
        input_variables=["draft_content"],
        template="""
        <draft_content>
        {draft_content}
        </draft_content>

        Please extract each cohesive section from the above draft. As a rule of thumb, please stick 
        to extracting text between H3 headings. If there are 2 or fewer H3 headings, then please use 
        H2 headings as delimiters instead.
        Your output should be a Python list where each item represents a single section of the draft.
        """
    )
    chain = section_extract_prompt | llm | StrOutputParser()
    sections = chain.invoke({"draft_content": draft_content})
    sections = sections.replace('[', '').replace(']', '')
    sections = sections.replace('`', '').replace('"', '').replace("'", '')
    sections = [i.strip() for i in sections.split(',')]
    return sections

def refine_section(section, seo_data):
    """Use LangChain to refine a single section."""
    print(f'section: {section}')
    print(f'seo_data: {seo_data}')
    final_seo_data = '\n\n'.join(seo_data)
    refine_prompt = PromptTemplate(
        input_variables=["final_seo_data","section"],
        template="""
        <SEO data>
        {final_seo_data}
        </SEO data>

        Based on the above SEO data, please expand and refine the below section into a well-developed paragraph or set of paragraphs.
        The word count for the finished, newly developed section should be around 100-200 words. 
        Ensure it is engaging, clear, and suitable for publication as part of an article. Do not include an introduction or conclusion.

        <section>
        {section}
        </section>
        """
    )

    chain = refine_prompt | llm | StrOutputParser()
    return chain.invoke({"section": section, "final_seo_data": final_seo_data})

# Submit button
if st.button("Generate SEO Strategy"):
    # Store input values into a dictionary for processing
    user_inputs = {
        "business_name": business_name,
        "business_description": business_description,
        "business_type": business_type,
        "niche": niche,
        "target_audience": target_audience,
        "keyword_focus": keyword_focus,
        "topic": topic,
    }

    try:
        # Step 1: Generate Foundation SEO Strategy
        foundation_chain = foundation_template | llm | StrOutputParser()
        foundation_content = foundation_chain.invoke(user_inputs)

        st.write("## Foundation SEO Strategy")
        st.write(foundation_content)

        # Step 2: Extract Pillar Page Ideas
        pillar_extraction_chain = {"foundation_content": foundation_chain} | pillar_extraction_template | llm | StrOutputParser()
        pillar_page_ideas = pillar_extraction_chain.invoke(user_inputs)
        pillar_page_ideas = pillar_page_ideas.replace('[', '').replace(']', '')
        pillar_page_ideas = pillar_page_ideas.replace('`', '').replace('"', '').replace("'", '')
        pillar_page_ideas = pillar_page_ideas.replace('python', '').replace('\n', '')
        pillar_page_ideas = [i.strip() for i in pillar_page_ideas.split(',')]
        st.write("## Extracted Pillar Page Ideas")
        st.write(pillar_page_ideas)

        # Step 3: Generate Keywords
        keywords_chain = {"foundation_content": foundation_chain} | keywords_template | llm | StrOutputParser()
        keywords_result = keywords_chain.invoke(user_inputs)
        st.write("## Keywords with Secondary Keywords")
        st.write(keywords_result)

        # Step 4: Generate Pillar Page Subtopics
        pillars_chain = {"foundation_content": foundation_chain} | pillars_template | llm | StrOutputParser()
        pillars_result = pillars_chain.invoke(user_inputs)
        st.write("## Pillar Pages with Subtopics")
        st.write(pillars_result)

        # Step 5: Generate Audience Questions
        questions_chain = {"foundation_content": foundation_chain} | questions_template | llm | StrOutputParser()
        questions_result = questions_chain.invoke(user_inputs)
        st.write("## Audience Questions")
        st.write(questions_result)

        # print(pillar_page_ideas)

        # Step 6: Loop over Pillar Page Ideas for Outlines using Threads
        outlines_only = []

        def generate_outline(index,idea):
            outline_chain = outline_template | llm | StrOutputParser()
            outline_dict = {
                "pillar_page_idea": idea,
                "foundation_content": foundation_content,
                "keywords": keywords_result,
                "questions": questions_result,
            }
            outline_result = outline_chain.invoke(outline_dict)
            outlines_only.append((index, outline_result))

        threads = []
        for i, idea in enumerate(pillar_page_ideas):
            thread = threading.Thread(target=generate_outline, args=(i,idea))
            threads.append(thread)
            thread.start()

        for thread in threads:
            thread.join()

        # Sort outlines by their original order
        outlines_only.sort(key=lambda x: x[0])
        outlines_only = [outline for _, outline in outlines_only]

        st.write("## Generated Outlines")
        for i, outline in enumerate(outlines_only):
            st.write(f"### Outline for Pillar Page Idea {i + 1}")
            st.write(outline)

        # Step 7: Loop over outlines to generate first drafts
        firstdrafts_only = []

        def generate_firstdraft(index,outline):
            firstdraft_chain = firstdraft_template | llm | StrOutputParser()
            firstdraft_dict = {
                "outline": outline,
                "business_name": foundation_content,
                "business_description": business_description,
                "features_page": features_page,
                "money_page": money_page,
            }
            firstdraft_result = firstdraft_chain.invoke(firstdraft_dict)
            firstdrafts_only.append((index, firstdraft_result))

        threads = []
        for i, outline in enumerate(outlines_only):
            thread = threading.Thread(target=generate_firstdraft, args=(i,outline))
            threads.append(thread)
            thread.start()

        for thread in threads:
            thread.join()

        # Sort outlines by their original order
        firstdrafts_only.sort(key=lambda x: x[0])
        firstdrafts_only = [draft for _, draft in firstdrafts_only]

        st.write("## Generated Drafts")
        for i, draft in enumerate(firstdrafts_only):
            st.write(f"### Draft for Pillar Page Idea {i + 1}")
            st.write(draft)

        # TODO: this articles step has really odd behavior. The outputs are shuffled 
        # and sometimes they seem to be talking about something else.
        # # Step 8: Loop over drafts to generate articles
        # articles_only = []
        # def refine_draft(idx, draft):
        #     """Process an entire draft, refining each section and combining them."""
        #     sections = extract_sections(draft)
        #     print(sections)
        #     refined_sections = []

        #     def refine_worker(section, index):
        #         seo_data = [foundation_content, '\n'.join(pillar_page_ideas), pillars_result, keywords_result, questions_result]
        #         refined_sections.append((index, refine_section(section, seo_data)))

        #     threads = []
        #     for i, section in enumerate(sections):
        #         thread = threading.Thread(target=refine_worker, args=(section, i))
        #         threads.append(thread)
        #         thread.start()

        #     for thread in threads:
        #         thread.join()

        #     # Sort sections by original order
        #     refined_sections.sort(key=lambda x: x[0])
        #     articles_only.append((idx, "\n\n".join([content for _, content in refined_sections])))
        #     return "\n\n".join([content for _, content in refined_sections])

        # threads = []
        # for i, draft in enumerate(firstdrafts_only):
        #     thread = threading.Thread(target=refine_draft, args=(i,draft))
        #     threads.append(thread)
        #     thread.start()

        # for thread in threads:
        #     thread.join()

        # # Sort articles by their original order
        # articles_only.sort(key=lambda x: x[0])
        # articles_only = [article for _, article in articles_only]

        # st.write("## Generated Articles")
        # for i, article in enumerate(articles_only):
        #     st.write(f"### Article for Pillar Page Idea {i + 1}")
        #     st.write(article)

        # Export Results
        all_results = {
            "Foundation SEO Strategy": foundation_content,
            "Pillar Page Ideas": "\n".join(pillar_page_ideas),
            "Keywords": keywords_result,
            "Pillar Subtopics": pillars_result,
            "Audience Questions": questions_result,
            "Outlines": "\n".join(outlines_only),
            "Drafts": "\n".join(firstdrafts_only),
            # "Articles": "\n".join(articles_only),
        }

        st.write("### Export Options")
        today_str = datetime.datetime.today().strftime('%Y-%m-%d')
        export_to_csv(all_results, f"{business_name}_seo_strategy_{today_str}.csv")
        export_to_word(all_results, f"{business_name}_seo_strategy_{today_str}.docx")

    except Exception as e:
        st.error(f"Error processing the prompts: {e}")
        error_ct += 1
        if error_ct > 100:
            print("Exiting due to high number of errors.")
            sys.exit()
